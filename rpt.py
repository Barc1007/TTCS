from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for section in doc.sections:
    section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
    section.left_margin=Cm(3); section.right_margin=Cm(2)

TNR=lambda size,bold=False: None
def H(doc,text,lvl=1,rgb=(0,51,102)):
    p=doc.add_heading(text,level=lvl)
    for r in p.runs: r.font.color.rgb=RGBColor(*rgb); r.font.name='Times New Roman'
    return p
def P(doc,text,bold=False,size=13):
    p=doc.add_paragraph(); run=p.add_run(text)
    run.font.name='Times New Roman'; run.font.size=Pt(size); run.bold=bold
    p.paragraph_format.space_after=Pt(6); p.paragraph_format.line_spacing=1.3
    return p
def B(doc,text):
    p=doc.add_paragraph(text,style='List Bullet')
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(13)
    p.paragraph_format.line_spacing=1.3; return p
def N(doc,text):
    p=doc.add_paragraph(text,style='List Number')
    for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(13)
    return p
def T(doc,headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style='Table Grid'
    h=t.rows[0].cells
    for i,v in enumerate(headers):
        h[i].text=v
        for para in h[i].paragraphs:
            for r in para.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(11)
    for ri,row in enumerate(rows):
        c=t.rows[ri+1].cells
        for ci,val in enumerate(row):
            c[ci].text=str(val)
            for para in c[ci].paragraphs:
                for r in para.runs: r.font.name='Times New Roman'; r.font.size=Pt(11)
    doc.add_paragraph()

# ===== TRANG BÌA =====
doc.add_paragraph()
for txt,sz,bold in [('HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG',14,True),('KHOA CÔNG NGHỆ THÔNG TIN',13,True)]:
    p=doc.add_paragraph(); p.alignment=1; r=p.add_run(txt)
    r.font.name='Times New Roman'; r.font.size=Pt(sz); r.bold=bold
for _ in range(5): doc.add_paragraph()
for txt,sz,rgb in [('BÁO CÁO CUỐI KỲ MÔN HỌC',22,None),('PHẦN MỀM QUẢN LÝ CƯ DÂN – ResidentIQ',17,(0,51,102))]:
    p=doc.add_paragraph(); p.alignment=1; r=p.add_run(txt)
    r.font.name='Times New Roman'; r.font.size=Pt(sz); r.bold=True
    if rgb: r.font.color.rgb=RGBColor(*rgb)
for _ in range(4): doc.add_paragraph()
for line in ['Nhóm: 17  |  Môn: Thiết kế Hệ thống Phần mềm','Thành viên:','  1. Nguyễn Thành Nam – N23DCCN108','  2. Nguyễn Hữu Đạt   – N23DCCN077','  3. Nguyễn Kỳ Đức An – N23DCCN070','','Hà Nội, Tháng 6/2026']:
    p=doc.add_paragraph(); p.alignment=1; r=p.add_run(line)
    r.font.name='Times New Roman'; r.font.size=Pt(13)
doc.add_page_break()

# ===== CHƯƠNG 1: YÊU CẦU (Tuần 1-3 tóm tắt) =====
H(doc,'CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG (Tuần 1–3)',1)
P(doc,'Đây là giai đoạn khảo sát và xác định toàn bộ yêu cầu cho hệ thống Quản lý Cư dân. Vấn đề cốt lõi: ban quản lý chung cư/nhà trọ đang dùng sổ sách giấy, dẫn đến khó tra cứu, dễ sai sót và không theo dõi được tạm trú/tạm vắng theo thời gian thực.')
H(doc,'1.1 Phạm vi & Người dùng',2)
T(doc,['Loại người dùng','Role trên hệ thống','Quyền hạn chính'],
[['Cán bộ hành chính','staff','Thêm/sửa hồ sơ cư dân, đăng ký tạm trú/tạm vắng'],
 ['Ban quản lý chung cư','admin','Toàn quyền: xóa cư dân, tạo tài khoản staff, xem báo cáo'],
 ['Cư dân','resident','Chỉ xem hồ sơ cá nhân và đổi mật khẩu']])
H(doc,'1.2 Yêu cầu chức năng chính (21 REQ)',2)
T(doc,['Nhóm','Mô tả'],
[['Xác thực (REQ-01→04)','Đăng nhập bằng username/email + password. Phân quyền 3 role: admin, staff, resident. Ghi nhận thời gian đăng nhập.'],
 ['Quản lý cư dân (REQ-05→11)','Thêm mới (CCCD unique, bắt buộc: name, cccd, dob, gender, room, regdate). Khi thêm tự động tạo tài khoản User liên kết. Sửa/xóa/tìm kiếm theo tên hoặc CCCD.'],
 ['Tạm trú/Tạm vắng (REQ-12→15)','Đăng ký tạm trú: ghi nhận address, start, end, reason, phone. Đăng ký tạm vắng: ghi nhận destination, start, end, reason, phone. Ràng buộc end >= start.'],
 ['Lịch sử biến động (REQ-16)','Mọi thao tác trên hồ sơ cư dân đều được ghi vào mảng history[] bên trong document Resident.'],
 ['Thống kê & Báo cáo (REQ-17→21)','Dashboard thống kê real-time: tổng cư dân, phân loại Thường trú/Tạm trú/Tạm vắng, biểu đồ xu hướng 6 tháng, log hoạt động gần nhất.']])
H(doc,'1.3 Yêu cầu phi chức năng (15 NFR)',2)
T(doc,['Nhóm NFR','Yêu cầu cụ thể'],
[['Bảo mật','Mật khẩu bắt buộc hash bằng bcrypt (salt=10). JWT Token thời hạn 7 ngày. Middleware requireAuth kiểm tra mọi route. Middleware requireRole kiểm tra quyền trước khi thực thi.'],
 ['Hiệu năng','API CRUD phản hồi < 500ms. Thao tác tìm kiếm trả kết quả < 1s.'],
 ['Khả dụng & Ràng buộc','Ứng dụng Web (SPA), không cài đặt thêm. Hỗ trợ Chrome, Firefox, Edge. Cở sở dữ liệu lưu trữ trên cloud (MongoDB Atlas).'],
 ['Dữ liệu','CCCD cư dân phải unique. Không xóa cư dân đang Tạm trú hoặc Tạm vắng (BR-006). Cư dân mới tạo tài khoản bắt buộc đổi mật khẩu lần đầu (mustChangePassword=true).']])

# ===== CHƯƠNG 2: KIẾN TRÚC (Tuần 4) =====
H(doc,'CHƯƠNG 2. KIẾN TRÚC HỆ THỐNG (Tuần 4)',1)
H(doc,'2.1 Architectural Drivers & Quyết định Stack',2)
P(doc,'Nhóm xác định 4 drivers chính dẫn dắt lựa chọn kiến trúc: (1) Team nhỏ 3 người - cần stack đồng nhất, (2) Dữ liệu linh hoạt có nested array (history, tamTru, tamVang) - cần NoSQL, (3) Tách Frontend/Backend rõ ràng để dễ mở rộng, (4) Triển khai nhanh trên cloud.')
P(doc,'Kiến trúc được chọn: Modular Monolith – một hệ thống server duy nhất nhưng có phân chia module rõ ràng theo chức năng (auth, residents, households, activityLog). Phù hợp với team nhỏ, deadline ngắn, tránh overhead của Microservices.')
H(doc,'2.2 Technology Stack thực tế',2)
T(doc,['Layer','Công nghệ','Version','Vai trò'],
[['Frontend','React.js + Vite','React 19, Vite 8','SPA, React-Router-DOM v7 cho điều hướng, Chart.js v4 cho biểu đồ Dashboard'],
 ['Backend','Node.js + Express.js','Express 5','RESTful API server, ES Module (type=module), chạy port 5000'],
 ['Database','MongoDB Atlas','Mongoose 8','Cloud NoSQL database, kết nối qua MONGODB_URI trong .env'],
 ['Auth & Security','jsonwebtoken + bcryptjs + helmet','JWT 9, bcrypt 3, helmet 8','JWT 7 ngày, bcrypt salt=10, helmet bảo vệ HTTP headers'],
 ['Validation','zod','Zod 4','Schema validation cho dữ liệu API đầu vào'],
 ['Dev Tools','nodemon + eslint','nodemon 3','Hot-reload development, code linting']])
H(doc,'2.3 Cấu hình môi trường thực tế',2)
P(doc,'Backend sử dụng file .env để tách biệt cấu hình nhạy cảm khỏi source code:')
T(doc,['Biến','Giá trị','Ý nghĩa'],
[['PORT','5000','Cổng lắng nghe của Express server'],
 ['MONGODB_URI','mongodb+srv://huudatnguyen1007_db_user:***@cluster0.sf32ppj.mongodb.net/?appName=Cluster0','Chuỗi kết nối MongoDB Atlas (cluster0.sf32ppj.mongodb.net)'],
 ['JWT_SECRET','ttcs2024-super-secret-jwt-key-xYz9Abc','Khóa bí mật ký JWT token'],
 ['CORS_ORIGIN','http://localhost:5173','Domain Frontend được phép gọi API']])
P(doc,'Frontend (.env): VITE_API_URL=http://localhost:5000/api – Base URL để gọi các API endpoints.')
H(doc,'2.4 Cấu trúc thư mục Backend',2)
T(doc,['Thư mục/File','Vai trò'],
[['src/server.js','Entry point: khởi động Express app, kết nối MongoDB Atlas, lắng nghe PORT'],
 ['src/app.js','Cấu hình Express: đăng ký middleware (cors, helmet, morgan, json), mount routes'],
 ['src/config/db.js','Hàm connectDB() gọi mongoose.connect()'],
 ['src/config/env.js','Đọc và export các biến môi trường (PORT, MONGODB_URI, JWT_SECRET, CORS_ORIGIN)'],
 ['src/models/','Mongoose Schemas: User, Resident, Household, ActivityLog'],
 ['src/controllers/','Business logic: auth, resident, household, activityLog controllers'],
 ['src/routes/','Express Routers: gắn controllers vào HTTP endpoints'],
 ['src/middleware/auth.middleware.js','requireAuth + requireRole – kiểm tra JWT và phân quyền'],
 ['src/utils/jwt.js','signToken() và verifyToken() dùng jsonwebtoken']])

# ===== CHƯƠNG 3: THIẾT KẾ CẤP CAO (Tuần 5) =====
H(doc,'CHƯƠNG 3. THIẾT KẾ CẤP CAO (Tuần 5)',1)
H(doc,'3.1 Thiết kế Thành phần (Component Design)',2)
P(doc,'Hệ thống theo mô hình MVC 3 tầng: Presentation (React SPA), Business Logic (Express Controllers), Data Access (Mongoose Models). Các module được phân tách rõ theo chức năng:')
T(doc,['Module','Thành phần','Trách nhiệm','REQ đáp ứng'],
[['Auth Module','AuthController, auth.routes.js','Đăng nhập, đăng ký (admin only), xem thông tin user, đổi mật khẩu','REQ-01→04'],
 ['Resident Module','ResidentController, resident.routes.js','CRUD cư dân, đăng ký tạm trú/vắng, thống kê, xem hồ sơ cá nhân','REQ-05→16'],
 ['Household Module','HouseholdController, household.routes.js','Danh sách và tạo mới hộ gia đình','Hỗ trợ REQ-05'],
 ['Activity Log Module','ActivityLogController, activityLog.routes.js','Truy xuất nhật ký hoạt động hệ thống','REQ-16'],
 ['Frontend SPA','React pages + Router','Hiển thị UI, gọi API, quản lý state đăng nhập qua AuthContext','Toàn bộ REQ']])
H(doc,'3.2 Thiết kế API (RESTful Endpoints)',2)
P(doc,'Toàn bộ API có prefix /api, sử dụng JSON. Xác thực qua header Authorization: Bearer <token>.')
T(doc,['Method','Endpoint','Chức năng','Quyền truy cập'],
[['POST','/api/auth/login','Đăng nhập. Body: {usernameOrEmail, password}. Trả về: {token, user}','Public'],
 ['POST','/api/auth/register','Tạo tài khoản staff/resident. Chỉ admin gọi được','admin'],
 ['GET','/api/auth/me','Lấy thông tin user đang đăng nhập từ token','Đã đăng nhập'],
 ['POST','/api/auth/change-password','Đổi mật khẩu. Body: {newPassword} (>=8 ký tự)','Đã đăng nhập'],
 ['GET','/api/residents','Danh sách tất cả cư dân, sắp xếp mới nhất','admin, staff'],
 ['POST','/api/residents','Thêm cư dân mới (tự tạo User tương ứng)','admin, staff'],
 ['GET','/api/residents/stats','Thống kê tổng hợp cho Dashboard','admin, staff'],
 ['GET','/api/residents/me','Cư dân xem hồ sơ cá nhân (qua residentId)','resident'],
 ['GET','/api/residents/:id','Chi tiết một cư dân theo MongoDB _id','Đã đăng nhập'],
 ['PUT','/api/residents/:id','Cập nhật thông tin cư dân','admin, staff'],
 ['DELETE','/api/residents/:id','Xóa cư dân (phải là Thường trú)','admin'],
 ['POST','/api/residents/:id/tam-tru','Đăng ký Tạm trú. Body: {address, start, end, reason, phone}','admin, staff'],
 ['POST','/api/residents/:id/tam-vang','Đăng ký Tạm vắng. Body: {destination, start, end, reason, phone}','admin, staff'],
 ['GET','/api/households','Danh sách hộ gia đình (populate headResidentId)','Đã đăng nhập'],
 ['POST','/api/households','Tạo hộ gia đình. Body: {code, apartment, address}','admin, staff']])
H(doc,'3.3 Thiết kế Dữ liệu – MongoDB Collections',2)
P(doc,'Hệ thống dùng 4 Collections trong MongoDB Atlas (cluster0.sf32ppj.mongodb.net):')
T(doc,['Collection','Trường dữ liệu chính','Ràng buộc quan trọng'],
[['users','name, username*, email*, passwordHash, role (admin/staff/resident), residentId (ref:Resident), mustChangePassword, timestamps','username và email: unique. role: enum. passwordHash không bao giờ trả về client.'],
 ['residents','name, cccd*, dob, gender, room, status (enum), address, ethnic, religion, job, relation, regdate, tamTru{}, tamVang{}, history[], timestamps','cccd: unique. status enum: Thường trú/Tạm trú/Tạm vắng. tamTru & tamVang lưu dạng embedded object.'],
 ['households','code*, apartment, address, headResidentId (ref:Resident), members[] (ref:Resident), timestamps','code: unique. Quan hệ 1-N với Resident qua ObjectId.'],
 ['activitylogs','actor, action, entityType, entityId, before{}, after{}, timestamps','Ghi nhận trạng thái trước và sau khi thay đổi (before/after).']])
P(doc,'(*) = unique. Quan hệ giữa users và residents: users.residentId tham chiếu tới residents._id (chỉ với role=resident).')
H(doc,'3.4 Thiết kế Bảo mật',2)
T(doc,['Cơ chế bảo mật','Cách triển khai thực tế'],
[['Mã hóa mật khẩu','bcrypt.hash(password, 10) – hash một chiều với cost factor 10. bcrypt.compare() để kiểm tra khi đăng nhập.'],
 ['JWT Authentication','jwt.sign({id, role, username}, JWT_SECRET, {expiresIn:"7d"}). Token được gửi qua header Authorization: Bearer <token>.'],
 ['Middleware requireAuth','Trích xuất Bearer token, jwt.verify() giải mã, User.findById() xác nhận user còn tồn tại. Gán req.user cho các handler sau.'],
 ['Middleware requireRole','requireRole("admin") hoặc requireRole("admin","staff") – kiểm tra req.user.role trước khi cho phép thực thi handler.'],
 ['CORS Protection','Chỉ chấp nhận request từ CORS_ORIGIN=http://localhost:5173. Các domain khác bị từ chối.'],
 ['HTTP Headers','helmet() tự động set các security headers: X-Content-Type-Options, X-Frame-Options, v.v...']])

# ===== CHUONG 4: THIET KE CAP THAP (Tuan 6) =====
H(doc,'CHÆ¯Æ NG 4. THIáº¾T Káº¾ Cáº¤P THáº¤P (Tuáº§n 6)',1)
H(doc,'4.1 SÆ¡ Ä‘á»“ Lá»›p â€“ Backend Classes',2)
T(doc,['Lá»›p','PhÆ°Æ¡ng thá»©c chÃ­nh','Phá»¥ thuá»™c'],
[['AuthController','register(), login(), me(), changePassword()','User model, bcrypt, jwt.js'],
 ['ResidentController','listResidents(), getResidentById(), createResident(), updateResident(), deleteResident(), registerTamTru(), registerTamVang(), getMyResidentInfo(), getResidentStats()','Resident model, User model, bcrypt'],
 ['HouseholdController','listHouseholds(), createHousehold()','Household model'],
 ['authMiddleware','requireAuth(req,res,next), requireRole(...roles)','User model, jwt.js'],
 ['jwt.js','signToken(payload): táº¡o token 7 ngÃ y, verifyToken(token): giáº£i mÃ£ token','jsonwebtoken, env.js'],
 ['User Model','Mongoose Schema: username unique, email unique, passwordHash, role enum, residentId ref Resident, mustChangePassword','mongoose'],
 ['Resident Model','Mongoose Schema: cccd unique, name, dob, gender, room, status enum, tamTru embedded, tamVang embedded, history array','mongoose'],
 ['Household Model','Mongoose Schema: code unique, apartment, address, headResidentId ref Resident, members array ref Resident','mongoose'],
 ['ActivityLog Model','Mongoose Schema: actor, action, entityType, entityId, before, after (Mixed type)','mongoose']])

H(doc,'4.2 SÆ¡ Ä‘á»“ TrÃ¬nh tá»± â€“ Luá»“ng ÄÄƒng nháº­p',2)
for s in [
  '1. Client goi POST /api/auth/login voi body {usernameOrEmail, password}.',
  '2. login() kiem tra body: thieu truong bat buoc thi tra ve 400 Bad Request.',
  '3. User.findOne de tim kiem user theo username hoac email trong MongoDB.',
  '4. Khong tim thay user -> tra 401: Thong tin dang nhap khong chinh xac.',
  '5. bcrypt.compare(password, user.passwordHash) kiem tra mat khau.',
  '6. Mat khau sai -> tra 401: Thong tin dang nhap khong chinh xac.',
  '7. Hop le: jwt.sign voi payload {id, role, username}, ky bang JWT_SECRET, expiresIn 7d.',
  '8. Tra 200 OK voi {token, user: sanitizeUser(user)} (loai bo truong passwordHash).',
]: N(doc,s)

H(doc,'4.3 SÆ¡ Ä‘á»“ TrÃ¬nh tá»± â€“ ThÃªm cÆ° dÃ¢n & Táº¡o tÃ i khoáº£n',2)
for s in [
  '1. POST /api/residents voi body day du cac truong bat buoc.',
  '2. requireAuth: jwt.verify() giai ma token, User.findById() xac nhan user ton tai.',
  '3. requireRole(admin, staff): kiem tra req.user.role -> khong du quyen tra 403.',
  '4. Kiem tra cac truong bat buoc: cccd, name, dob, gender, room, regdate -> thieu tra 400.',
  '5. Resident.findOne kiem tra CCCD da ton tai chua -> co roi tra 409 Conflict.',
  '6. Resident.create: luu ho so voi status mac dinh Thuong tru, history ghi Them moi cu dan.',
  '7. defaultPassword = 8 so cuoi CCCD: cccd.slice(-8).',
  '8. bcrypt.hash ma hoa mat khau mac dinh voi salt rounds = 10.',
  '9. User.create: username=cccd, email=cccd@resident.local, role=resident, mustChangePassword=true, residentId lien ket.',
  '10. Tra 201 Created voi {message, resident}.',
]: N(doc,s)

H(doc,'4.4 SÆ¡ Ä‘á»“ TrÃ¬nh tá»± â€“ ÄÄƒng kÃ½ Táº¡m trÃº',2)
for s in [
  '1. POST /api/residents/:id/tam-tru voi body {address, start, end, reason, phone}.',
  '2. requireAuth + requireRole(admin, staff) kiem tra token va quyen.',
  '3. Resident.findById(id) -> khong tim thay tra 404.',
  '4. Thieu address, start, end hoac phone -> tra 400: Thieu thong tin tam tru.',
  '5. So sanh end < start -> tra 400: Ngay ket thuc khong duoc nho hon ngay bat dau.',
  '6. resident.status dang la Tam vang -> tra 409: Cu dan dang tam vang, khong the dang ky tam tru.',
  '7. Cap nhat: resident.status = Tam tru; resident.tamTru = {address, start, end, reason, phone}.',
  '8. addHistory: push {action: Dang ky tam tru, by: ten can bo, at: thoi gian hien tai} vao history[].',
  '9. resident.save() luu thay doi xuong MongoDB Atlas.',
  '10. Tra 200 OK voi {message, resident da cap nhat}.',
]: N(doc,s)

H(doc,'4.5 Thiáº¿t káº¿ Frontend â€“ MÃ n hÃ¬nh & Routing',2)
P(doc,'Frontend la SPA React dieu huong bang React-Router-DOM v7. App.jsx xu ly phan quyen UI dua tren user.role:')
T(doc,['Route','Component','Quyen','Chuc nang'],
[['/dashboard','Dashboard.jsx','admin, staff','Thong ke: tong so cu dan, phan loai trang thai, bieu do Line Chart 6 thang (Chart.js), log hoat dong gan nhat'],
 ['/residents','ResidentList.jsx','admin, staff','Danh sach cu dan, tim kiem realtime, loc theo trang thai'],
 ['/residents/add','ResidentAdd.jsx','admin, staff','Form them cu dan moi co validation phia client'],
 ['/residents/:id','ResidentDetail.jsx','admin, staff','Xem chi tiet ho so, sua thong tin, nut Dang ky tam tru/tam vang, xem lich su bien dong'],
 ['/tamtru','TamTru.jsx','admin, staff','Form dang ky tam tru cho cu dan duoc chon'],
 ['/tamvang','TamVang.jsx','admin, staff','Form dang ky tam vang cho cu dan duoc chon'],
 ['/reports','Reports.jsx','admin, staff','Trang bao cao va thong ke tong hop'],
 ['/my-profile','MyProfile.jsx','resident','Cu dan xem ho so ca nhan (goi GET /api/residents/me)']])
P(doc,'Hai Context Provider quan ly state toan app: AuthContext (user, token, login, logout) va ResidentContext (danh sach cu dan, toast, modal xac nhan).')

# ===== CHUONG 5: KE HOACH TRIEN KHAI & KIEM THU (Tuan 7) =====
H(doc,'CHÆ¯Æ NG 5. Káº¾ HOáº CH TRIá»‚N KHAI VÃ€ KIá»‚M THá»¬ (Tuáº§n 7)',1)

H(doc,'5.1 Thiáº¿t káº¿ Thuáº­t toÃ¡n â€“ Kiá»ƒm tra Business Rules',2)
P(doc,'Há»‡ thá»‘ng cÃ i Ä‘áº·t cÃ¡c Business Rule trá»±c tiáº¿p trong ResidentController. DÆ°á»›i Ä‘Ã¢y lÃ  2 thuáº­t toÃ¡n quan trá»ng nháº¥t:')
P(doc,'Thuáº­t toÃ¡n xÃ³a cÆ° dÃ¢n (BR-006 â€“ KhÃ´ng xÃ³a khi Ä‘ang lÆ°u trÃº):', bold=True)
for s in [
  'Äáº¦U VÃ€O: id cÆ° dÃ¢n cáº§n xÃ³a.',
  'BÆ°á»›c 1: TÃ¬m cÆ° dÃ¢n theo id. Náº¿u khÃ´ng tá»“n táº¡i -> tráº£ 404.',
  'BÆ°á»›c 2: Kiá»ƒm tra resident.status == "Táº¡m trÃº" -> tráº£ 409: KhÃ´ng thá»ƒ xÃ³a cÆ° dÃ¢n Ä‘ang táº¡m trÃº.',
  'BÆ°á»›c 3: Kiá»ƒm tra resident.status == "Táº¡m váº¯ng" -> tráº£ 409: KhÃ´ng thá»ƒ xÃ³a cÆ° dÃ¢n Ä‘ang táº¡m váº¯ng.',
  'BÆ°á»›c 4: User.deleteOne vá»›i residentId = resident._id Ä‘á»ƒ xÃ³a tÃ i khoáº£n liÃªn káº¿t.',
  'BÆ°á»›c 5: resident.deleteOne() xÃ³a há»“ sÆ¡ cÆ° dÃ¢n.',
  'Äáº¦U RA: 200 OK - "XÃ³a cÆ° dÃ¢n thÃ nh cÃ´ng".',
]: B(doc,s)

P(doc,'Thuáº­t toÃ¡n tÃ­nh thá»‘ng kÃª Dashboard (getResidentStats):', bold=True)
for s in [
  'Äáº¦U VÃ€O: KhÃ´ng cÃ³ tham sá»‘ (láº¥y toÃ n bá»™).',
  'BÆ°á»›c 1: Resident.find() láº¥y táº¥t cáº£ cÆ° dÃ¢n (chá»‰ láº¥y trÆ°á»ng status, createdAt, history, name).',
  'BÆ°á»›c 2: Äáº¿m total, thuongtru, tamtru, tamvang báº±ng Array.filter().',
  'BÆ°á»›c 3: VÃ²ng láº·p 6 thÃ¡ng gáº§n nháº¥t: vá»›i má»—i thÃ¡ng Ä‘áº¿m sá»‘ cÆ° dÃ¢n cÃ³ createdAt <= endOfMonth.',
  'BÆ°á»›c 4: Gá»™p táº¥t cáº£ history tá»« má»i cÆ° dÃ¢n, sáº¯p xáº¿p giáº£m dáº§n theo thá»i gian, láº¥y 10 má»¥c Ä‘áº§u.',
  'Äáº¦U RA: {total, thuongtru, tamtru, tamvang, monthlyStats[6], recentActivities[10]}.',
]: B(doc,s)

H(doc,'5.2 Xá»­ lÃ½ lá»—i (Error Handling Strategy)',2)
T(doc,['HTTP Code','TrÆ°á»ng há»£p thá»±c táº¿ trong Source Code'],
[['400 Bad Request','Thiáº¿u trÆ°á»ng báº¯t buá»™c (cccd, name, dob...). NgÃ y káº¿t thÃºc < ngÃ y báº¯t Ä‘áº§u (táº¡m trÃº/táº¡m váº¯ng). Máº­t kháº©u má»›i Ã­t hÆ¡n 8 kÃ½ tá»±.'],
 ['401 Unauthorized','Sai username/password khi Ä‘Äƒng nháº­p. KhÃ´ng cÃ³ Bearer token. Token khÃ´ng há»£p lá»‡ hoáº·c háº¿t háº¡n (7 ngÃ y). User Ä‘Ã£ bá»‹ xÃ³a khá»i DB nhÆ°ng token cÃ²n hiá»‡u lá»±c.'],
 ['403 Forbidden','Resident cá»‘ gá»i API chá»‰ dÃ nh cho staff/admin. Staff cá»‘ gá»i API DELETE chá»‰ dÃ nh cho admin.'],
 ['404 Not Found','id cÆ° dÃ¢n khÃ´ng tá»“n táº¡i trong MongoDB. TÃ i khoáº£n chÆ°a liÃªn káº¿t residentId khi gá»i /residents/me.'],
 ['409 Conflict','ThÃªm cÆ° dÃ¢n cÃ³ CCCD Ä‘Ã£ tá»“n táº¡i. ÄÄƒng kÃ½ táº¡m trÃº khi Ä‘ang táº¡m váº¯ng (vÃ  ngÆ°á»£c láº¡i). XÃ³a cÆ° dÃ¢n Ä‘ang Táº¡m trÃº hoáº·c Táº¡m váº¯ng. MÃ£ há»™ gia Ä‘Ã¬nh Ä‘Ã£ tá»“n táº¡i.'],
 ['500 Internal Error','Lá»—i káº¿t ná»‘i MongoDB Atlas. Lá»—i Mongoose schema validation. Lá»—i ngoáº¡i lá»‡ khÃ´ng Ä‘Æ°á»£c báº¯t.']])
P(doc,'Táº¥t cáº£ API tráº£ vá» lá»—i theo chuáº©n JSON: { "message": "<mÃ´ táº£ lá»—i>" }. HÃ m sanitizeUser() Ä‘áº£m báº£o trÆ°á»ng passwordHash khÃ´ng bao giá» xuáº¥t hiá»‡n trong response.')

H(doc,'5.3 Chiáº¿n lÆ°á»£c Kiá»ƒm thá»­',2)
T(doc,['Loáº¡i kiá»ƒm thá»­','CÃ´ng cá»¥','Pháº¡m vi & Ká»‹ch báº£n kiá»ƒm thá»­'],
[['Integration Testing (API)','Postman','Kiá»ƒm thá»­ 15 endpoints. Test case positive: Ä‘Äƒng nháº­p Ä‘Ãºng -> nháº­n token. Test case negative: Ä‘Äƒng nháº­p sai -> 401. Táº¡o 2 cÆ° dÃ¢n trÃ¹ng CCCD -> 409. ÄÄƒng kÃ½ táº¡m trÃº khi Ä‘ang táº¡m váº¯ng -> 409. XÃ³a cÆ° dÃ¢n Ä‘ang Táº¡m trÃº -> 409. Gá»i API khÃ´ng cÃ³ token -> 401.'],
 ['Manual Testing (UI)','TrÃ¬nh duyá»‡t Chrome','ÄÄƒng nháº­p vá»›i tá»«ng role (admin, staff, resident), kiá»ƒm tra quyá»n háº¡n hiá»ƒn thá»‹ menu. ThÃªm cÆ° dÃ¢n -> kiá»ƒm tra tÃ i khoáº£n tá»± táº¡o. ÄÄƒng nháº­p báº±ng CCCD lÃ m username vá»›i máº­t kháº©u 8 sá»‘ cuá»‘i. Báº¯t buá»™c Ä‘á»•i máº­t kháº©u láº§n Ä‘áº§u (mustChangePassword). Dashboard: kiá»ƒm tra sá»‘ liá»‡u vÃ  biá»ƒu Ä‘á»“ cáº­p nháº­t Ä‘Ãºng sau khi thÃªm cÆ° dÃ¢n.'],
 ['Data Validation','Trá»±c tiáº¿p trÃªn UI','Nháº­p form thiáº¿u trÆ°á»ng -> kiá»ƒm tra thÃ´ng bÃ¡o lá»—i validation. Nháº­p ngÃ y táº¡m trÃº end < start -> kiá»ƒm tra API tráº£ 400.']])

H(doc,'5.4 TiÃªu chuáº©n Code (Coding Standards)',2)
T(doc,['Quy táº¯c','Chi tiáº¿t Ã¡p dá»¥ng trong Source Code'],
[['Äáº·t tÃªn hÃ m/biáº¿n','camelCase: getResidentStats, createResident, addHistory, signToken, verifyToken'],
 ['Äáº·t tÃªn Model/Class','PascalCase: Resident, User, Household, ActivityLog, AuthController'],
 ['Äáº·t tÃªn file','Dáº¡ng <module>.<type>.js: auth.controller.js, resident.routes.js, auth.middleware.js'],
 ['Module System','ES Modules (import/export) cho cáº£ Backend vÃ  Frontend (package.json: "type":"module")'],
 ['Cáº¥u trÃºc thÆ° má»¥c Backend','src/config (db, env), src/models, src/controllers, src/routes, src/middleware, src/utils'],
 ['Comment','Ghi chÃº Business Rules trong code: BR-006 (khong xoa khi tam tru/tam vang), ghi chÃº logic route (/me phai truoc /:id)'],
 ['Error handling','Moi controller dung try/catch. Middleware xu ly loi tap trung. Phan biet ro HTTP codes 400/401/403/404/409/500']])

H(doc,'5.5 Káº¿ hoáº¡ch Triá»ƒn khai',2)
T(doc,['Giai Ä‘oáº¡n','Nhiá»‡m vá»¥','Káº¿t quáº£'],
[['Giai Ä‘oáº¡n 1 (Tuan 1-4)','Phan tich yeu cau, thiet ke kien truc, chon MERN Stack, cau hinh MongoDB Atlas, xay dung Mongoose Schemas (User, Resident, Household, ActivityLog)','DB Schema hoan chinh, ket noi Atlas thanh cong'],
 ['Giai Ä‘oáº¡n 2 (Tuan 5-6)','Trien khai day du Backend API (auth, residents, households), Middleware, JWT Auth. Xay dung Frontend React: tat ca pages, contexts, routing, Chart.js Dashboard','Backend API chay o port 5000, Frontend chay o port 5173'],
 ['Giai Ä‘oáº¡n 3 (Tuan 7)','Kiem thu toan bo chuc nang (Postman + Manual). Sua loi. Kiem tra phan quyen tung role. Dam bao cac Business Rules hoat dong dung.','He thong on dinh, day du chuc nang theo yeu cau']])

# ===== KET LUAN =====
H(doc,'Káº¾T LUáº¬N',1)
P(doc,'Dá»± Ã¡n Pháº§n má»m Quáº£n lÃ½ CÆ° dÃ¢n â€“ ResidentIQ Ä‘Ã£ Ä‘Æ°á»£c thiáº¿t káº¿ vÃ  triá»ƒn khai hoÃ n chá»‰nh báº±ng MERN Stack (MongoDB Atlas, Express.js, React.js, Node.js). Há»‡ thá»‘ng Ä‘Ã¡p á»©ng Ä‘áº§y Ä‘á»§ 21 yÃªu cáº§u chá»©c nÄƒng vÃ  15 yÃªu cáº§u phi chá»©c nÄƒng Ä‘á» ra tá»« Tuáº§n 1.')
P(doc,'Äiá»ƒm ná»•i báº­t ká»¹ thuáº­t cá»§a dá»± Ã¡n: (1) Kiáº¿n trÃºc Modular Monolith phÃ¢n chia module rÃµ rÃ ng phÃ¹ há»£p quy mÃ´ team. (2) Há»‡ thá»‘ng phÃ¢n quyá»n 3 cáº¥p (admin/staff/resident) Ä‘Æ°á»£c báº£o vá»‡ á»Ÿ cáº£ Backend (middleware) vÃ  Frontend (route guard). (3) TÃ­nh nÄƒng tá»± Ä‘á»™ng táº¡o tÃ i khoáº£n User khi thÃªm cÆ° dÃ¢n. (4) Lá»‹ch sá»­ biáº¿n Ä‘á»™ng Ä‘Æ°á»£c ghi nháº­n tá»± Ä‘á»™ng má»i thao tÃ¡c. (5) CÃ¡c Business Rules nhÆ° chá»‘ng xÃ³a khi Ä‘ang lÆ°u trÃº (BR-006), CCCD duy nháº¥t (BR-001), chá»‘ng táº¡m trÃº khi Ä‘ang táº¡m váº¯ng (BR-005) Ä‘Æ°á»£c cÃ i Ä‘áº·t cháº·t cháº½ trong controller.')
P(doc,'BÃ¡o cÃ¡o nÃ y pháº£n Ã¡nh chÃ­nh xÃ¡c toÃ n bá»™ quÃ¡ trÃ¬nh tá»« phÃ¢n tÃ­ch yÃªu cáº§u Ä‘áº¿n thiáº¿t káº¿ kiáº¿n trÃºc vÃ  cÃ i Ä‘áº·t thá»±c táº¿ trong source code, phá»¥c vá»¥ má»¥c tiÃªu Ä‘Ã¡nh giÃ¡ cuá»‘i ká»³ mÃ´n Thiáº¿t káº¿ Há»‡ thá»‘ng Pháº§n má»m.')

out = r'd:\TTCS-main\BaoCaoCuoiKy_Final.docx'
doc.save(out)
print('DONE:', out)
