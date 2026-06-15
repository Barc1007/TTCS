from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

def add_heading(doc, text, level=1, color=(0,51,102)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Times New Roman'
    return p

def add_para(doc, text, bold=False, size=13):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    doc.add_paragraph()

# ===== TRANG BIA =====
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
r2.font.name = 'Times New Roman'; r2.font.size = Pt(13); r2.bold = True

for _ in range(4): doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title.add_run('BÁO CÁO CUỐI KỲ MÔN HỌC')
rt.font.name = 'Times New Roman'; rt.font.size = Pt(20); rt.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = subtitle.add_run('PHẦN MỀM QUẢN LÝ CƯ DÂN')
rs.font.name = 'Times New Roman'; rs.font.size = Pt(18); rs.bold = True
rs.font.color.rgb = RGBColor(0, 51, 102)

for _ in range(3): doc.add_paragraph()

info_lines = [
    'Nhóm 17 – Môn Thiết kế hệ thống phần mềm',
    '',
    'Thành viên:',
    '  1. Nguyễn Thành Nam – N23DCCN108',
    '  2. Nguyễn Hữu Đạt   – N23DCCN077',
    '  3. Nguyễn Kỳ Đức An – N23DCCN070',
    '',
    'Hà Nội, tháng 6 năm 2026',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

doc.add_page_break()

# ===== TỔNG QUAN =====
add_heading(doc, 'TỔNG QUAN DỰ ÁN', 1)
add_para(doc, 'Phần mềm Quản lý Cư dân được xây dựng nhằm hỗ trợ ban quản lý chung cư/khu dân cư quản lý thông tin nhân khẩu một cách hiệu quả, thay thế phương pháp ghi sổ tay truyền thống dễ sai sót. Báo cáo này tóm tắt quá trình thiết kế và phát triển dự án từ Tuần 1 đến Tuần 7, phản ánh đúng kiến trúc và công nghệ đã sử dụng trong Source Code thực tế.', size=13)

add_heading(doc, '1. GIAI ĐOẠN 1 – XÁC ĐỊNH YÊU CẦU (Tuần 1–3)', 1)

add_heading(doc, '1.1 Tầm nhìn, Phạm vi & Đối tượng (Tuần 1-2)', 2)
add_para(doc, 'Dự án tập trung vào việc quản lý cư dân, hộ gia đình, tạm trú, tạm vắng và thống kê biến động nhân khẩu. Hệ thống phân chia thành 3 vai trò người dùng chính: Cán bộ quản lý (staff), Ban quản lý (admin) và Cư dân (resident).', size=13)

add_heading(doc, '1.2 Yêu cầu chức năng & phi chức năng (Tuần 3)', 2)
add_table(doc,
    ['Nhóm', 'Yêu cầu tiêu biểu'],
    [
        ['Xác thực & phân quyền', 'Đăng nhập bằng username/password; phân quyền theo vai trò (admin, staff, resident)'],
        ['Quản lý cư dân', 'Thêm mới (tự động cấp tài khoản), cập nhật, xóa, tìm kiếm theo CCCD/Tên'],
        ['Tạm trú / Tạm vắng', 'Đăng ký tạm trú, tạm vắng; đảm bảo ràng buộc thời gian (end >= start)'],
        ['Thống kê & Báo cáo', 'Thống kê tổng quan số lượng cư dân theo trạng thái cư trú'],
        ['Bảo mật (NFR)', 'Mã hóa mật khẩu bằng thuật toán bcrypt, phân quyền truy cập thông qua JWT'],
        ['Hiệu năng (NFR)', 'Hệ thống thiết kế theo kiến trúc Single Page Application (SPA) giúp tải trang nhanh'],
    ]
)

# ===== KIẾN TRÚC =====
add_heading(doc, '2. KIẾN TRÚC HỆ THỐNG VÀ CÔNG NGHỆ (Tuần 4)', 1)

add_heading(doc, '2.1 Architectural Drivers', 2)
add_para(doc, 'Các quyết định kiến trúc được định hướng bởi các yếu tố sau:', size=13)
add_table(doc,
    ['Driver', 'Tác động kiến trúc'],
    [
        ['Năng lực đội ngũ & Thời gian', 'Chọn 스택 (Stack) phổ biến, dễ phát triển nhanh: MERN Stack.'],
        ['Dữ liệu phi cấu trúc', 'Sử dụng cơ sở dữ liệu Document-based (NoSQL) để lưu các mảng lịch sử (history) linh hoạt.'],
        ['Bảo mật thông tin cá nhân', 'Xác thực qua JWT, mật khẩu hash bcrypt, thiết kế middleware Auth chặt chẽ.'],
        ['Dễ dàng triển khai (Cloud)', 'Sử dụng dịch vụ Cloud Database (MongoDB Atlas) tích hợp kết nối URI trực tiếp.'],
    ]
)

add_heading(doc, '2.2 Technology Stack Thực Tế', 2)
add_table(doc,
    ['Thành phần', 'Công nghệ', 'Lý do & Vai trò'],
    [
        ['Frontend', 'React.js + Vite', 'Tạo giao diện Single Page Application (SPA), render dữ liệu nhanh, quản lý state linh hoạt.'],
        ['Backend', 'Node.js + Express.js', 'Xây dựng RESTful API xử lý bất đồng bộ tốt, phù hợp phát triển tốc độ cao.'],
        ['Database', 'MongoDB Atlas', 'Lưu trữ dữ liệu phi cấu trúc (Document), dùng Mongoose để định nghĩa Schema.'],
        ['Authentication', 'jsonwebtoken & bcryptjs', 'Tạo JWT Token stateless và mã hóa mật khẩu an toàn.'],
    ]
)

# ===== THIẾT KẾ CẤP CAO =====
add_heading(doc, '3. THIẾT KẾ CẤP CAO (Tuần 5)', 1)

add_heading(doc, '3.1 Thiết kế Cơ sở dữ liệu (MongoDB)', 2)
add_para(doc, 'Sử dụng MongoDB với Mongoose Schema, đảm bảo ràng buộc tính toàn vẹn dữ liệu.', size=13)
add_table(doc,
    ['Collection', 'Trường dữ liệu chính', 'Ràng buộc & Ghi chú'],
    [
        ['users', 'name, username, email, passwordHash, role, residentId, mustChangePassword', 'username và email là unique. role enum: admin, staff, resident.'],
        ['residents', 'name, cccd, dob, gender, room, status, address, ethnic, regdate, tamTru, tamVang, history[]', 'cccd là unique. Lịch sử (history) lưu lại các thao tác thay đổi.'],
        ['households', 'code, apartment, address, headResidentId, members[]', 'code là unique. headResidentId liên kết với collection residents.'],
        ['activitylogs', 'actor, action, entityType, entityId, before, after', 'Lưu vết (audit log) chi tiết mọi hoạt động trong hệ thống.'],
    ]
)

add_heading(doc, '3.2 Cấu hình Môi trường & Triển khai', 2)
add_para(doc, 'Hệ thống sử dụng file môi trường (.env) để quản lý cấu hình kết nối, bảo mật thông tin nhạy cảm. Cấu hình Backend:', size=13)
add_table(doc,
    ['Biến môi trường', 'Giá trị', 'Ý nghĩa'],
    [
        ['PORT', '5000', 'Cổng chạy server Node.js backend'],
        ['MONGODB_URI', 'mongodb+srv://huudatnguyen1007_db_user:[PASS]@cluster0.sf32ppj.mongodb.net/?appName=Cluster0', 'Chuỗi kết nối trực tiếp lên Database Cloud MongoDB Atlas'],
        ['CORS_ORIGIN', 'http://localhost:5173', 'Cấp quyền truy cập cho Frontend (React/Vite) chạy ở port 5173'],
    ]
)

# ===== THIẾT KẾ CẤP THẤP =====
add_heading(doc, '4. THIẾT KẾ CẤP THẤP (Tuần 6)', 1)

add_heading(doc, '4.1 Thiết kế Lớp (Controllers & Models)', 2)
add_table(doc,
    ['Controller/Lớp', 'Phương thức / Trách nhiệm', 'Model liên kết'],
    [
        ['AuthController', 'register(), login(), me(), changePassword() - Xử lý xác thực người dùng.', 'User'],
        ['ResidentController', 'listResidents(), getResidentById(), createResident(), updateResident(), deleteResident(), registerTamTru(), registerTamVang(), getResidentStats()', 'Resident, User'],
        ['HouseholdController', 'listHouseholds(), createHousehold() - Xử lý tạo và quản lý hộ gia đình.', 'Household'],
        ['ActivityLogController', 'Ghi nhận và truy xuất log hệ thống.', 'ActivityLog'],
    ]
)

add_heading(doc, '4.2 Thiết kế API (RESTful)', 2)
add_table(doc,
    ['Method', 'Endpoint', 'Chức năng', 'Auth Middleware'],
    [
        ['POST', '/api/auth/login', 'Đăng nhập vào hệ thống', 'Không'],
        ['GET', '/api/residents', 'Lấy danh sách cư dân', 'Có (verifyToken)'],
        ['POST', '/api/residents', 'Thêm cư dân & cấp tài khoản tự động', 'Có (verifyToken)'],
        ['PUT', '/api/residents/:id', 'Cập nhật thông tin cư dân', 'Có (verifyToken)'],
        ['POST', '/api/residents/:id/tamtru', 'Đăng ký trạng thái Tạm trú', 'Có (verifyToken)'],
        ['GET', '/api/residents/stats', 'Thống kê số liệu Dashboard', 'Có (verifyToken)'],
    ]
)

# ===== TRIỂN KHAI & LUỒNG XỬ LÝ =====
add_heading(doc, '5. LẬP KẾ HOẠCH TRIỂN KHAI (Tuần 7)', 1)

add_heading(doc, '5.1 Luồng xử lý Thêm cư dân & Tạo tài khoản (Thuật toán)', 2)
add_para(doc, 'Khi thêm mới cư dân, hệ thống không chỉ tạo hồ sơ mà còn tự động sinh tài khoản đăng nhập cho cư dân đó:', size=13)
seq_steps = [
    '1. Nhận yêu cầu POST /api/residents với dữ liệu (cccd, name, dob, ...)',
    '2. Kiểm tra ràng buộc: cccd đã tồn tại trong Resident DB chưa?',
    '3. Nếu chưa, tạo bản ghi Resident mới với trạng thái mặc định "Thường trú".',
    '4. Tự động sinh mật khẩu mặc định (ví dụ: 8 số cuối CCCD).',
    '5. Hash mật khẩu bằng bcrypt.hash().',
    '6. Tạo bản ghi User mới với role = "resident", username = cccd, residentId = [ID vừa tạo], mustChangePassword = true.',
    '7. Trả về Response 201 thành công.',
]
for s in seq_steps:
    p = doc.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

add_heading(doc, '5.2 Luồng xử lý Đăng ký Tạm trú/Tạm vắng', 2)
algo = [
    '1. Xác thực ID cư dân từ URL (req.params.id).',
    '2. Kiểm tra thời gian: Ngày kết thúc (end) phải >= Ngày bắt đầu (start).',
    '3. Kiểm tra logic trạng thái: Không cho đăng ký Tạm trú nếu đang Tạm vắng (và ngược lại).',
    '4. Cập nhật field tamTru hoặc tamVang và đổi field status của cư dân.',
    '5. Ghi nhận log vào mảng resident.history (Lịch sử biến động).',
    '6. Lưu thay đổi xuống Database.',
]
for s in algo:
    p = doc.add_paragraph(s, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

add_heading(doc, 'KẾT LUẬN', 1)
add_para(doc, 'Dự án "Phần mềm Quản lý Cư dân" đã được thiết kế và triển khai hoàn chỉnh bằng công nghệ MERN Stack. Các module cốt lõi như Xác thực người dùng, Quản lý cư dân, Đăng ký lưu trú (Tạm trú/Tạm vắng) và Bảng điều khiển thống kê (Dashboard) đã hoạt động trơn tru. Hệ thống sử dụng MongoDB Atlas cho lưu trữ đám mây, kết nối API qua Express.js, đảm bảo được các yêu cầu về tốc độ, bảo mật và sự tiện dụng. Báo cáo này đã chắt lọc nội dung toàn bộ 7 tuần thiết kế để làm rõ mối liên hệ giữa lý thuyết kiến trúc và mã nguồn thực tế.', size=13)

out = r'd:\TTCS-main\BaoCaoCuoiKy_Nhom17_V2.docx'
doc.save(out)
print('DONE:', out)
