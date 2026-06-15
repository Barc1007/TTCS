from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# Define styles
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

def add_heading(doc, text, level=1, color=(0,51,102)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Times New Roman'
    return p

def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    doc.add_paragraph()

# ==========================================
# TRANG BÌA
# ==========================================
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG')
r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
r2.font.name = 'Times New Roman'; r2.font.size = Pt(14); r2.bold = True

for _ in range(6): doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title.add_run('BÁO CÁO CUỐI KỲ MÔN HỌC')
rt.font.name = 'Times New Roman'; rt.font.size = Pt(24); rt.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = subtitle.add_run('THIẾT KẾ HỆ THỐNG PHẦN MỀM QUẢN LÝ CƯ DÂN')
rs.font.name = 'Times New Roman'; rs.font.size = Pt(20); rs.bold = True
rs.font.color.rgb = RGBColor(0, 51, 102)

for _ in range(5): doc.add_paragraph()

info_lines = [
    'Giảng viên hướng dẫn: .......................................',
    '',
    'Nhóm thực hiện: 17',
    'Thành viên:',
    '  1. Nguyễn Thành Nam – N23DCCN108',
    '  2. Nguyễn Hữu Đạt   – N23DCCN077',
    '  3. Nguyễn Kỳ Đức An – N23DCCN070',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True if 'Giảng viên' in line or 'Nhóm' in line else False

for _ in range(3): doc.add_paragraph()
p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_date = p_date.add_run('Hà Nội, Tháng 6/2026')
r_date.font.name = 'Times New Roman'
r_date.font.size = Pt(14)
r_date.italic = True

doc.add_page_break()

# ==========================================
# LỜI MỞ ĐẦU
# ==========================================
add_heading(doc, 'LỜI MỞ ĐẦU', 1)
add_para(doc, 'Báo cáo này trình bày toàn bộ quy trình từ khâu phân tích yêu cầu, thiết kế kiến trúc, thiết kế cấp cao và cấp thấp, cho đến kế hoạch triển khai của phần mềm "Quản lý Cư dân". Đây là báo cáo tổng hợp kết quả công việc từ Tuần 1 đến Tuần 7, được đúc kết ngắn gọn để thể hiện rõ ràng và chính xác các quyết định kỹ thuật tương ứng với mã nguồn thực tế đang được xây dựng (MERN Stack).')
add_para(doc, 'Phần mềm được phát triển nhằm mục tiêu chuyển đổi số quy trình quản lý hành chính tại các khu chung cư và nhà trọ, giảm thiểu sai sót do ghi chép thủ công, đồng thời tăng tốc độ xử lý các nghiệp vụ như tra cứu nhân khẩu, đăng ký tạm trú, tạm vắng và thống kê dân số trực quan.')
add_para(doc, 'Cấu trúc báo cáo bao gồm:')
add_bullet(doc, 'Chương 1: Phân tích yêu cầu hệ thống (Tóm tắt Tuần 1-3)')
add_bullet(doc, 'Chương 2: Thiết kế Kiến trúc và Công nghệ (Tuần 4)')
add_bullet(doc, 'Chương 3: Thiết kế Cấp cao (Tuần 5)')
add_bullet(doc, 'Chương 4: Thiết kế Cấp thấp (Tuần 6)')
add_bullet(doc, 'Chương 5: Kế hoạch triển khai và Kiểm thử (Tuần 7)')

# ==========================================
# CHƯƠNG 1
# ==========================================
add_heading(doc, 'CHƯƠNG 1. PHÂN TÍCH YÊU CẦU HỆ THỐNG (Tuần 1-3)', 1)

add_heading(doc, '1.1 Tầm nhìn và Phạm vi dự án', 2)
add_para(doc, 'Dự án giải quyết trực tiếp những bất cập trong việc quản lý nhân khẩu theo cách truyền thống. Các rủi ro như mất mát dữ liệu, khó kiểm soát thời hạn tạm trú/tạm vắng, và sự thiếu đồng bộ giữa các sổ sách được loại bỏ bằng một nền tảng tập trung duy nhất.')
add_para(doc, 'Phạm vi in-scope bao gồm: Quản lý chi tiết hồ sơ cư dân (thêm, sửa, xóa, tìm kiếm), quản lý hộ gia đình, theo dõi lịch sử biến động cư trú, và xuất các thống kê tổng quan (Dashboard). Phạm vi out-of-scope: Tích hợp hệ thống dữ liệu quốc gia về dân cư, nhận diện sinh trắc học, quản lý phí chung cư.')

add_heading(doc, '1.2 Phân tích Người dùng (User Personas)', 2)
add_table(doc, 
    ['Loại người dùng', 'Vai trò trên hệ thống', 'Mục tiêu chính'],
    [
        ['Cán bộ hành chính / Lễ tân', 'Staff', 'Nhập liệu hồ sơ cư dân, cập nhật thông tin, thực hiện thủ tục đăng ký tạm trú, tạm vắng hàng ngày.'],
        ['Ban quản lý chung cư', 'Admin', 'Giám sát toàn bộ dữ liệu, xem báo cáo tổng hợp, thống kê biến động dân cư để báo cáo chính quyền.'],
        ['Cư dân', 'Resident', 'Xem và theo dõi trạng thái cư trú của bản thân, xác nhận lại các thông tin cá nhân.']
    ]
)

add_heading(doc, '1.3 Danh sách Yêu cầu Chức năng (Functional Requirements)', 2)
add_para(doc, 'Hệ thống xác định các chức năng cốt lõi bắt buộc phải có, bao gồm:')
add_bullet(doc, 'REQ-01 (Xác thực): Đăng nhập hệ thống phân quyền chặt chẽ theo role Admin, Staff và Resident.')
add_bullet(doc, 'REQ-02 (Quản lý Cư dân): Cho phép Staff/Admin thêm mới cư dân. Khi thêm mới, hệ thống tự động sinh tài khoản User (username là CCCD) và cấp mật khẩu ngẫu nhiên/mặc định.')
add_bullet(doc, 'REQ-03 (Tạm trú/Tạm vắng): Cho phép cư dân chuyển trạng thái sang tạm trú hoặc tạm vắng. Ràng buộc thời gian bắt đầu phải nhỏ hơn hoặc bằng thời gian kết thúc.')
add_bullet(doc, 'REQ-04 (Lịch sử biến động): Mọi thao tác làm thay đổi dữ liệu của cư dân đều được ghi log lưu lại thời gian và người thực hiện.')
add_bullet(doc, 'REQ-05 (Dashboard): Thống kê theo thời gian thực tổng số cư dân, phân loại thường trú, tạm trú, tạm vắng và biểu đồ biến động theo các tháng.')

add_heading(doc, '1.4 Danh sách Yêu cầu Phi Chức năng (Non-Functional Requirements)', 2)
add_bullet(doc, 'NFR-01 (Bảo mật): Mật khẩu người dùng tuyệt đối không lưu dạng plain-text mà phải được băm (hash) bằng thuật toán bcrypt. Sử dụng JWT (JSON Web Token) để xác thực các request API.')
add_bullet(doc, 'NFR-02 (Hiệu năng): Các API CRUD cơ bản phải phản hồi trong thời gian dưới 500ms. Thao tác lọc, tìm kiếm danh sách trên giao diện xử lý dưới 1 giây.')
add_bullet(doc, 'NFR-03 (Khả dụng): Hệ thống xây dựng dưới dạng ứng dụng Web (SPA), hỗ trợ tốt trên các trình duyệt hiện đại, không yêu cầu cài đặt phần mềm phụ.')

# ==========================================
# CHƯƠNG 2
# ==========================================
add_heading(doc, 'CHƯƠNG 2. KIẾN TRÚC HỆ THỐNG VÀ CÔNG NGHỆ (Tuần 4)', 1)

add_heading(doc, '2.1 Các yếu tố thúc đẩy kiến trúc (Architectural Drivers)', 2)
add_para(doc, 'Quyết định kiến trúc của phần mềm được dẫn dắt bởi 4 yếu tố chính (Drivers):')
add_bullet(doc, 'Thời gian phát triển nhanh (Time-to-market): Đội ngũ chỉ có 3 thành viên với giới hạn thời gian thực hiện ngắn. Do đó, lựa chọn một stack công nghệ đồng nhất (JavaScript cho cả Frontend và Backend) là ưu tiên số một để tối ưu hóa nguồn lực.')
add_bullet(doc, 'Đặc thù Dữ liệu phi cấu trúc (Unstructured Data): Các bản ghi tạm trú, tạm vắng, hay lịch sử biến động có thể thay đổi cấu trúc linh hoạt. Cơ sở dữ liệu NoSQL tỏ ra vượt trội hơn so với RDBMS truyền thống trong tình huống này.')
add_bullet(doc, 'Tách biệt Giao diện và Nghiệp vụ: Cần xây dựng API chuẩn RESTful độc lập hoàn toàn với giao diện để dễ dàng nâng cấp hoặc viết thêm ứng dụng Mobile sau này (NFR-03).')
add_bullet(doc, 'Khả năng triển khai Cloud dễ dàng: Hệ thống ưu tiên các giải pháp PaaS/DBaaS để giảm thiểu gánh nặng vận hành server vật lý.')

add_heading(doc, '2.2 Lựa chọn Technology Stack thực tế', 2)
add_para(doc, 'Dựa trên các Drivers trên, nhóm quyết định sử dụng cấu trúc MERN Stack (MongoDB, Express.js, React.js, Node.js). Cụ thể:')
add_table(doc,
    ['Thành phần', 'Công nghệ áp dụng', 'Lý do & Vai trò'],
    [
        ['Frontend', 'React.js (v19) + Vite', 'Tạo SPA mượt mà. Sử dụng React-Router cho việc điều hướng và Chart.js để vẽ biểu đồ Dashboard trực quan.'],
        ['Backend', 'Node.js + Express.js', 'Chạy môi trường runtime bất đồng bộ, tốc độ xử lý I/O cao. Express.js cung cấp bộ router RESTful mạnh mẽ.'],
        ['Database', 'MongoDB Atlas', 'Giải pháp Database-as-a-Service, lưu trữ dữ liệu dạng Document (JSON-like), linh hoạt mở rộng các trường Schema.'],
        ['Security & Auth', 'JWT, bcryptjs, helmet', 'Mã hóa mật khẩu, quản lý phiên đăng nhập stateless, helmet bảo vệ các HTTP headers khỏi tấn công XSS.'],
    ]
)

add_heading(doc, '2.3 Mô hình Triển khai & Cấu hình Môi trường', 2)
add_para(doc, 'Ứng dụng được cấu hình để dễ dàng triển khai với các biến môi trường (.env) phân định rõ ràng. Cấu hình Backend hiện hành:')
add_para(doc, 'PORT=5000\nMONGODB_URI=mongodb+srv://huudatnguyen1007_db_user:sTcS95NbaovvORDH@cluster0.sf32ppj.mongodb.net/?appName=Cluster0\nCORS_ORIGIN=http://localhost:5173', italic=True)
add_para(doc, 'Trong đó, chuỗi kết nối MONGODB_URI dẫn trực tiếp lên cụm cluster trên MongoDB Atlas. Cấu hình CORS đảm bảo chỉ Frontend đang chạy ở cổng 5173 mới được phép gọi API, chặn các domain lạ.')

# ==========================================
# CHƯƠNG 3
# ==========================================
add_heading(doc, 'CHƯƠNG 3. THIẾT KẾ CẤP CAO (Tuần 5)', 1)

add_heading(doc, '3.1 Thiết kế Thành phần (Component Diagram)', 2)
add_para(doc, 'Hệ thống tuân thủ chặt chẽ nguyên lý MVC (Model-View-Controller) cho Backend và Component-Based cho Frontend:')
add_bullet(doc, 'Presentation Layer (Frontend): Bao gồm các màn hình chính như LoginPage, DashboardPage, ResidentListPage. Tầng này chịu trách nhiệm hiển thị và gọi API (fetch/axios) gửi dữ liệu JSON lên Server.')
add_bullet(doc, 'Business Logic Layer (Backend Controllers): Bao gồm AuthController (xử lý đăng nhập), ResidentController (xử lý nghiệp vụ quản lý cư dân), HouseholdController (quản lý hộ gia đình). Tầng này validate input, kiểm tra business rule trước khi gọi DB.')
add_bullet(doc, 'Data Access Layer (Models): Giao tiếp với MongoDB qua thư viện Mongoose. Các Mongoose Schemas (User, Resident, Household) định nghĩa cấu trúc dữ liệu và các Hooks (pre-save) nếu cần.')

add_heading(doc, '3.2 Đặc tả API (RESTful Endpoints)', 2)
add_para(doc, 'Backend cung cấp một bộ API chuẩn hóa. Dưới đây là danh sách các Endpoint trọng yếu:')

add_table(doc,
    ['Method', 'Endpoint', 'Mục đích', 'Yêu cầu Token'],
    [
        ['POST', '/api/auth/login', 'Nhận username, password. Trả về JWT Token', 'Không'],
        ['GET', '/api/residents', 'Truy xuất toàn bộ danh sách cư dân', 'Có (Admin/Staff)'],
        ['POST', '/api/residents', 'Tạo hồ sơ cư dân mới, tự động tạo tài khoản User', 'Có (Admin/Staff)'],
        ['GET', '/api/residents/:id', 'Lấy thông tin chi tiết một cư dân theo ID', 'Có'],
        ['PUT', '/api/residents/:id', 'Cập nhật hồ sơ cư dân (họ tên, giới tính...)', 'Có (Admin/Staff)'],
        ['DELETE', '/api/residents/:id', 'Xóa cư dân (yêu cầu không đang tạm trú/tạm vắng)', 'Có (Admin/Staff)'],
        ['POST', '/api/residents/:id/tamtru', 'Khai báo thông tin đăng ký Tạm trú', 'Có (Admin/Staff)'],
        ['POST', '/api/residents/:id/tamvang', 'Khai báo thông tin đăng ký Tạm vắng', 'Có (Admin/Staff)'],
        ['GET', '/api/residents/stats', 'Lấy dữ liệu (Total, Biểu đồ tháng) cho Dashboard', 'Có'],
        ['GET', '/api/households', 'Lấy danh sách các hộ gia đình đang quản lý', 'Có'],
    ]
)

add_heading(doc, '3.3 Thiết kế Dữ liệu Cấp cao (Data Model)', 2)
add_para(doc, 'Vì sử dụng MongoDB, dữ liệu không chia thành nhiều bảng Normalize như SQL mà tận dụng khả năng Embed (nhúng dữ liệu) cho các trường hợp 1-nhiều ít thay đổi.')
add_bullet(doc, 'Thay vì tạo bảng riêng cho Tạm Trú, Tạm Vắng, chúng được lưu trực tiếp dưới dạng Object (tamTru, tamVang) trong bản ghi Resident.')
add_bullet(doc, 'Lịch sử biến động (history) được lưu thành một mảng (Array) các object bên trong Resident, giúp truy xuất nhanh chóng toàn bộ quá trình biến động của cá nhân đó chỉ với 1 query.')

# ==========================================
# CHƯƠNG 4
# ==========================================
add_heading(doc, 'CHƯƠNG 4. THIẾT KẾ CẤP THẤP (Tuần 6)', 1)

add_heading(doc, '4.1 Thiết kế Schema Cơ sở dữ liệu vật lý (Mongoose Schemas)', 2)
add_para(doc, 'Mỗi collection trong MongoDB được kiểm soát chặt chẽ bằng Mongoose Schema. Chi tiết các trường dữ liệu:')

add_para(doc, 'Bảng Users (Tài khoản người dùng):', bold=True)
add_bullet(doc, 'username: String, required, unique (thường lấy theo CCCD cư dân).')
add_bullet(doc, 'email: String, required, unique.')
add_bullet(doc, 'passwordHash: String (mật khẩu đã bị băm bởi bcrypt).')
add_bullet(doc, 'role: Enum [\'admin\', \'staff\', \'resident\'], mặc định là staff.')
add_bullet(doc, 'residentId: ObjectId (ref: \'Resident\'), null đối với Admin/Staff.')

add_para(doc, 'Bảng Residents (Hồ sơ cư dân):', bold=True)
add_bullet(doc, 'name, dob, gender, room: String, required.')
add_bullet(doc, 'cccd: String, required, unique (Đảm bảo quy tắc BR-001: 1 CCCD chỉ có 1 hồ sơ).')
add_bullet(doc, 'status: Enum [\'Thường trú\', \'Tạm trú\', \'Tạm vắng\'], default: \'Thường trú\'.')
add_bullet(doc, 'tamTru / tamVang: Object chứa { address/destination, start, end, reason, phone }.')
add_bullet(doc, 'history: Mảng các Object { action, by, at }. Tự động push khi có thao tác update.')

add_para(doc, 'Bảng ActivityLog (Log hệ thống độc lập):', bold=True)
add_bullet(doc, 'actor: String (Người thực hiện). action: String (Loại hành động). entityType & entityId: Tham chiếu tới bản ghi bị tác động.')

add_heading(doc, '4.2 Thiết kế Lớp (Controllers & Middleware)', 2)
add_para(doc, 'Chi tiết các hàm xử lý bên trong Backend:')
add_bullet(doc, 'Auth.Controller.js: Chứa hàm login(). Đọc username từ request, tìm trong DB. Dùng bcrypt.compare() so sánh mật khẩu. Nếu khớp, dùng jwt.sign() tạo mã token với payload là { id, role, residentId } trả về cho Frontend.')
add_bullet(doc, 'Resident.Controller.js: Hàm createResident() chịu trách nhiệm tạo mới dữ liệu. Sau khi lưu Resident.create(), nó tự động trigger tạo User.create() tương ứng. Hàm getResidentStats() xử lý vòng lặp tính toán tổng cư dân, phân rã theo 6 tháng gần nhất để trả về JSON cho Chart.js vẽ biểu đồ.')
add_bullet(doc, 'AuthMiddleware.js: Hàm verifyToken(). Trích xuất token từ header "Authorization: Bearer <token>". Dùng jwt.verify() để giải mã. Nếu hợp lệ, gán user vào req.user và gọi next(); nếu không hợp lệ trả về 401 Unauthorized.')

add_heading(doc, '4.3 Sơ đồ Trình tự (Sequence Flow) - Đăng ký Tạm trú', 2)
add_para(doc, 'Quy trình chi tiết khi một cán bộ thực hiện đăng ký tạm trú cho cư dân:')
seq_flow = [
    'Bước 1: Frontend gửi POST request tới /api/residents/:id/tamtru kèm Token và body (start, end, address, reason).',
    'Bước 2: Router chạy qua authMiddleware kiểm tra Token hợp lệ.',
    'Bước 3: ResidentController nhận request, gọi Resident.findById(id).',
    'Bước 4: Kiểm tra điều kiện thời gian (end phải >= start). Trả lỗi 400 nếu sai.',
    'Bước 5: Kiểm tra status hiện tại. Nếu resident.status == "Tạm vắng", trả lỗi 409 (Không thể tạm trú khi đang tạm vắng).',
    'Bước 6: Gán resident.status = "Tạm trú" và gán dữ liệu vào resident.tamTru.',
    'Bước 7: Gọi hàm helper addHistory() đẩy dòng "Đăng ký tạm trú" vào mảng history.',
    'Bước 8: Gọi resident.save() lưu xuống MongoDB.',
    'Bước 9: Controller trả về mã 200 OK kèm theo đối tượng Resident vừa cập nhật cho Frontend cập nhật UI.'
]
for s in seq_flow:
    p = doc.add_paragraph(s, style='List Number')
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

# ==========================================
# CHƯƠNG 5
# ==========================================
add_heading(doc, 'CHƯƠNG 5. KẾ HOẠCH TRIỂN KHAI VÀ KIỂM THỬ (Tuần 7)', 1)

add_heading(doc, '5.1 Thiết kế Thuật toán tạo tài khoản tự động', 2)
add_para(doc, 'Để giảm bớt quy trình cấp phát mật khẩu thủ công, hệ thống áp dụng thuật toán sinh tài khoản liền mạch trong quá trình nhập liệu:')
algo2 = [
    'Khởi tạo resident mới với dữ liệu từ req.body.',
    'Bóc tách 8 ký tự cuối cùng của cccd làm defaultPassword.',
    'Sử dụng bcrypt.hash(defaultPassword, salt=10) để mã hóa thành chuỗi hash.',
    'Khởi tạo user mới: username = cccd, email = cccd + "@resident.local", passwordHash = hash.',
    'Đặt role = "resident" và cờ mustChangePassword = true để ép đổi mật khẩu lần đầu.',
    'Thực hiện Lưu song song (Promise.all) hoặc tuần tự xuống Database.'
]
for s in algo2:
    add_bullet(doc, s)

add_heading(doc, '5.2 Chiến lược Xử lý lỗi (Error Handling)', 2)
add_para(doc, 'Mọi API trả về lỗi theo một chuẩn JSON thống nhất: { "message": "<Mô tả lỗi cụ thể>" }. Các HTTP Status Code được quy ước chặt chẽ:')
add_table(doc,
    ['HTTP Code', 'Ý nghĩa', 'Trường hợp áp dụng thực tế'],
    [
        ['400 Bad Request', 'Lỗi Input từ người dùng', 'Nhập thiếu trường bắt buộc, ngày kết thúc nhỏ hơn ngày bắt đầu.'],
        ['401 Unauthorized', 'Lỗi xác thực', 'Sai username/password, Token hết hạn, không truyền Token.'],
        ['403 Forbidden', 'Lỗi phân quyền', 'Resident cố gắng truy cập API xóa cư dân (chỉ dành cho Admin).'],
        ['404 Not Found', 'Không thấy dữ liệu', 'ID cư dân không tồn tại trong DB khi xem chi tiết.'],
        ['409 Conflict', 'Xung đột dữ liệu', 'Thêm cư dân có CCCD đã tồn tại. Xóa cư dân đang Tạm trú/Tạm vắng.'],
        ['500 Internal Server Error', 'Lỗi hệ thống', 'Mất kết nối MongoDB Atlas, lỗi try-catch ngoại lệ.']
    ]
)

add_heading(doc, '5.3 Kế hoạch Kiểm thử (Testing Strategy)', 2)
add_para(doc, 'Chất lượng phần mềm được đảm bảo thông qua 2 tầng kiểm thử:')
add_bullet(doc, 'Integration Testing: Sử dụng Postman để mô phỏng gọi API thực tế. Kiểm tra các trường hợp Negative: cố tình truyền thiếu JWT Token, cố tình tạo 2 cư dân trùng CCCD, cố tình đăng ký Tạm vắng khi đang Tạm trú để xem hệ thống có bắt lỗi 409 chính xác không.')
add_bullet(doc, 'Manual Testing (UI/UX): Đóng vai cán bộ quản lý (staff), đăng nhập vào giao diện React. Thực hiện quy trình từ Thêm cư dân -> Tìm kiếm -> Xem chi tiết -> Đăng ký tạm trú -> Qua trang Dashboard xem số lượng có tăng lên ngay lập tức (Real-time update) hay không.')

add_heading(doc, '5.4 Tiêu chuẩn Code (Coding Standards)', 2)
add_para(doc, 'Để duy trì sự thống nhất cho mã nguồn, team áp dụng các chuẩn sau:')
add_bullet(doc, 'Tên biến và hàm (Variables/Functions): Sử dụng camelCase. Ví dụ: getResidentStats, createResident.')
add_bullet(doc, 'Tên Model/Schema: Sử dụng PascalCase số ít. Ví dụ: Resident, User, Household.')
add_bullet(doc, 'Tên File: Sử dụng đuôi mô tả chức năng. Ví dụ: auth.controller.js, resident.routes.js.')
add_bullet(doc, 'Cấu trúc thư mục: Chia module chuẩn (src/models, src/controllers, src/routes, src/utils). Mọi nghiệp vụ đặt tại Controller, Model chỉ định nghĩa Schema.')

# ==========================================
# KẾT LUẬN
# ==========================================
add_heading(doc, 'KẾT LUẬN', 1)
add_para(doc, 'Dự án "Phần mềm Quản lý Cư dân" đã được thiết kế và triển khai một cách bài bản qua 7 tuần. Từ khâu xác định yêu cầu, vẽ ra cấu trúc dữ liệu cho tới việc lựa chọn Stack công nghệ MERN (MongoDB, Express, React, Node.js) phù hợp với nguồn lực. Điểm nổi bật của hệ thống là khả năng kiểm soát chặt chẽ các ràng buộc nghiệp vụ (tạm trú, tạm vắng, chống xóa khi đang lưu trú) và lưu vết lịch sử đầy đủ. Hệ thống API thiết kế chuẩn REST, dễ dàng mở rộng, tích hợp tốt với MongoDB Atlas trên nền tảng Cloud. Báo cáo này đã thể hiện một cái nhìn toàn cảnh, từ ý tưởng lý thuyết đến những dòng code thực tế đang vận hành hệ thống.')

out = r'd:\TTCS-main\BaoCaoCuoiKy_Nhom17_V3.docx'
doc.save(out)
print('DONE:', out)
