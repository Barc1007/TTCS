# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2)
    s.left_margin=Cm(3); s.right_margin=Cm(2)

def H(doc, text, lvl=1):
    p = doc.add_heading(text, level=lvl)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        if lvl == 1: r.font.color.rgb = RGBColor(0, 51, 102)
    return p

def P(doc, text, bold=False, sz=13):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = bold
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    return p

def B(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    return p

def N(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    return p

def T(doc, hdrs, rows, fsz=11):
    tbl = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    tbl.style = 'Table Grid'
    for i, h in enumerate(hdrs):
        c = tbl.rows[0].cells[i]; c.text = h
        for para in c.paragraphs:
            for r in para.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(fsz)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = str(v)
            for para in c.paragraphs:
                for r in para.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(fsz)
    doc.add_paragraph()

# ========== TRANG BÌA ==========
for txt, sz, bold in [
    ('HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG', 14, True),
    ('KHOA CÔNG NGHỆ THÔNG TIN', 13, True)
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = bold

for _ in range(5): doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('BÁO CÁO CUỐI KỲ MÔN HỌC')
r.font.name = 'Times New Roman'; r.font.size = Pt(22); r.bold = True

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẦN MỀM QUẢN LÝ CƯ DÂN – ResidentIQ')
r.font.name = 'Times New Roman'; r.font.size = Pt(17); r.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

for _ in range(4): doc.add_paragraph()

for line in ['Nhóm: 17  |  Môn: Thiết kế Hệ thống Phần mềm',
             'Thành viên:',
             '   1. Nguyễn Thành Nam – N23DCCN108',
             '   2. Nguyễn Hữu Đạt   – N23DCCN077',
             '   3. Nguyễn Kỳ Đức An – N23DCCN070',
             '', 'Hà Nội, Tháng 6/2026']:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line); r.font.name = 'Times New Roman'; r.font.size = Pt(13)

doc.add_page_break()

# ========== CHƯƠNG 1 ==========
H(doc, 'CHƯƠNG 1. PHÂN TÍCH YÊU CẦU (Tuần 1–3)', 1)
P(doc, 'Giai đoạn này xác định toàn bộ yêu cầu cho hệ thống. Vấn đề cốt lõi: ban quản lý chung cư/nhà trọ dùng sổ sách giấy dẫn đến khó tra cứu, dễ sai sót và không theo dõi được tạm trú/tạm vắng theo thời gian thực.')

H(doc, '1.1 Người dùng hệ thống', 2)
T(doc, ['Người dùng', 'Role', 'Quyền hạn'],
[['Cán bộ hành chính', 'staff', 'Thêm/sửa hồ sơ cư dân, đăng ký tạm trú/tạm vắng'],
 ['Ban quản lý', 'admin', 'Toàn quyền: xóa cư dân, tạo tài khoản, xem báo cáo'],
 ['Cư dân', 'resident', 'Chỉ xem hồ sơ cá nhân và đổi mật khẩu']])

H(doc, '1.2 Yêu cầu chức năng & phi chức năng', 2)
T(doc, ['Nhóm', 'Yêu cầu'],
[['Xác thực (REQ-01→04)', 'Đăng nhập username/email + password. Phân quyền 3 role. Đổi mật khẩu bắt buộc lần đầu (mustChangePassword).'],
 ['Quản lý cư dân (REQ-05→11)', 'CRUD cư dân. CCCD bắt buộc unique. Khi thêm cư dân tự động tạo tài khoản User liên kết. Tìm kiếm theo tên/CCCD.'],
 ['Tạm trú/Tạm vắng (REQ-12→15)', 'Đăng ký tạm trú: {address, start, end, reason, phone}. Đăng ký tạm vắng: {destination, start, end, reason, phone}. Ràng buộc end >= start.'],
 ['Biến động & Thống kê (REQ-16→21)', 'Ghi lịch sử mọi thao tác vào mảng history[]. Dashboard: tổng cư dân, phân loại 3 trạng thái, biểu đồ 6 tháng, log 10 hoạt động gần nhất.'],
 ['Bảo mật (NFR)', 'Mật khẩu hash bcrypt salt=10. JWT 7 ngày. Middleware requireAuth + requireRole bảo vệ mọi route. CORS chỉ cho phép localhost:5173.'],
 ['Hiệu năng & Khả dụng (NFR)', 'API CRUD phản hồi <500ms. SPA không cần cài đặt. Chạy tốt trên Chrome, Firefox, Edge.']])

# ========== CHƯƠNG 2 ==========
H(doc, 'CHƯƠNG 2. KIẾN TRÚC HỆ THỐNG (Tuần 4)', 1)
P(doc, 'Kiến trúc được chọn: Modular Monolith – một server duy nhất, phân chia rõ theo module chức năng. Phù hợp team 3 người, deadline ngắn, tránh overhead của Microservices.')

H(doc, '2.1 Technology Stack', 2)
T(doc, ['Layer', 'Công nghệ', 'Version', 'Vai trò'],
[['Frontend', 'React.js + Vite', 'React 19, Vite 8', 'SPA, React-Router-DOM v7, Chart.js v4 cho Dashboard'],
 ['Backend', 'Node.js + Express.js', 'Express 5', 'RESTful API, ES Module (type=module), port 5000'],
 ['Database', 'MongoDB Atlas', 'Mongoose 8', 'Cloud NoSQL, kết nối qua MONGODB_URI'],
 ['Auth & Security', 'jsonwebtoken + bcryptjs + helmet', 'JWT 9, bcrypt 3', 'Token 7 ngày, hash salt=10, HTTP headers bảo vệ'],
 ['Validation', 'zod', 'Zod 4', 'Schema validation đầu vào API']])

H(doc, '2.2 Cấu hình môi trường (.env)', 2)
T(doc, ['Biến', 'Giá trị', 'Ý nghĩa'],
[['PORT', '5000', 'Cổng Express server'],
 ['MONGODB_URI', 'mongodb+srv://huudatnguyen1007_db_user:***@cluster0.sf32ppj.mongodb.net', 'Kết nối MongoDB Atlas'],
 ['JWT_SECRET', 'ttcs2024-super-secret-jwt-key-xYz9Abc', 'Khóa ký JWT'],
 ['CORS_ORIGIN', 'http://localhost:5173', 'Domain Frontend được phép gọi API']])
P(doc, 'Frontend: VITE_API_URL=http://localhost:5000/api')

H(doc, '2.3 Cấu trúc thư mục Backend', 2)
T(doc, ['Thư mục/File', 'Vai trò'],
[['src/server.js', 'Entry point: kết nối MongoDB Atlas, lắng nghe PORT'],
 ['src/config/db.js + env.js', 'connectDB() và đọc biến môi trường'],
 ['src/models/', 'Mongoose Schemas: User, Resident, Household, ActivityLog'],
 ['src/controllers/', 'Business logic: auth, resident, household, activityLog'],
 ['src/routes/', 'Express Routers gắn controllers vào HTTP endpoints'],
 ['src/middleware/auth.middleware.js', 'requireAuth (JWT verify) + requireRole (kiểm tra quyền)'],
 ['src/utils/jwt.js', 'signToken() – tạo JWT 7 ngày | verifyToken() – giải mã']])

# ========== CHƯƠNG 3 ==========
H(doc, 'CHƯƠNG 3. THIẾT KẾ CẤP CAO (Tuần 5)', 1)

H(doc, '3.1 Thiết kế API – RESTful Endpoints', 2)
T(doc, ['Method', 'Endpoint', 'Chức năng', 'Quyền'],
[['POST', '/api/auth/login', 'Đăng nhập → trả {token, user}', 'Public'],
 ['POST', '/api/auth/register', 'Tạo tài khoản staff/resident', 'admin'],
 ['GET', '/api/auth/me', 'Thông tin user hiện tại', 'Đăng nhập'],
 ['POST', '/api/auth/change-password', 'Đổi mật khẩu (>=8 ký tự)', 'Đăng nhập'],
 ['GET', '/api/residents', 'Danh sách tất cả cư dân', 'admin, staff'],
 ['POST', '/api/residents', 'Thêm cư dân + tự tạo User', 'admin, staff'],
 ['GET', '/api/residents/stats', 'Thống kê Dashboard', 'admin, staff'],
 ['GET', '/api/residents/me', 'Cư dân xem hồ sơ cá nhân', 'resident'],
 ['GET', '/api/residents/:id', 'Chi tiết một cư dân', 'Đăng nhập'],
 ['PUT', '/api/residents/:id', 'Cập nhật thông tin', 'admin, staff'],
 ['DELETE', '/api/residents/:id', 'Xóa cư dân (chỉ Thường trú)', 'admin'],
 ['POST', '/api/residents/:id/tam-tru', 'Đăng ký Tạm trú', 'admin, staff'],
 ['POST', '/api/residents/:id/tam-vang', 'Đăng ký Tạm vắng', 'admin, staff'],
 ['GET/POST', '/api/households', 'DS hộ gia đình / Tạo mới', 'Đăng nhập']])

H(doc, '3.2 Thiết kế MongoDB Collections', 2)
T(doc, ['Collection', 'Trường dữ liệu chính', 'Ràng buộc'],
[['users', 'name, username*, email*, passwordHash, role (admin/staff/resident), residentId (ref Resident), mustChangePassword', 'username & email: unique. passwordHash không trả về client.'],
 ['residents', 'name, cccd*, dob, gender, room, status (enum), address, ethnic, regdate, tamTru{address,start,end,reason,phone}, tamVang{destination,start,end,reason,phone}, history[]', 'cccd: unique. status enum: Thường trú/Tạm trú/Tạm vắng.'],
 ['households', 'code*, apartment, address, headResidentId (ref Resident), members[] (ref Resident)', 'code: unique. 1-N với Resident.'],
 ['activitylogs', 'actor, action, entityType, entityId, before{}, after{}', 'Ghi trạng thái trước/sau khi thay đổi.']])

H(doc, '3.3 Thiết kế Bảo mật', 2)
T(doc, ['Cơ chế', 'Triển khai thực tế'],
[['Hash mật khẩu', 'bcrypt.hash(password, 10) khi lưu. bcrypt.compare() khi đăng nhập.'],
 ['JWT Token', 'jwt.sign({id, role, username}, JWT_SECRET, {expiresIn:"7d"}). Gửi qua header Authorization: Bearer <token>.'],
 ['requireAuth', 'Trích Bearer token → jwt.verify() → User.findById() xác nhận tồn tại → gán req.user.'],
 ['requireRole', 'Kiểm tra req.user.role trong danh sách cho phép → không đủ quyền trả 403.'],
 ['Helmet', 'Tự động set security headers: X-Content-Type-Options, X-Frame-Options, v.v...']])

# ========== CHƯƠNG 4 ==========
H(doc, 'CHƯƠNG 4. THIẾT KẾ CẤP THẤP (Tuần 6)', 1)

H(doc, '4.1 Sơ đồ Lớp (Class Diagram)', 2)
T(doc, ['Lớp', 'Phương thức', 'Phụ thuộc'],
[['AuthController', 'register(), login(), me(), changePassword()', 'User, bcryptjs, jwt.js'],
 ['ResidentController', 'listResidents(), getResidentById(), createResident(), updateResident(), deleteResident(), registerTamTru(), registerTamVang(), getMyResidentInfo(), getResidentStats()', 'Resident, User, bcryptjs'],
 ['HouseholdController', 'listHouseholds(), createHousehold()', 'Household'],
 ['authMiddleware', 'requireAuth(req,res,next), requireRole(...roles)', 'User, jwt.js'],
 ['jwt.js', 'signToken(payload), verifyToken(token)', 'jsonwebtoken, env.js']])

H(doc, '4.2 Sơ đồ Trình tự – Đăng nhập', 2)
for s in [
    '1. Client → POST /api/auth/login {usernameOrEmail, password}',
    '2. login(): kiểm tra body đủ trường → thiếu trả 400.',
    '3. User.findOne({$or:[{username},{email}]}) tìm trong MongoDB.',
    '4. Không tìm thấy → 401 "Thông tin đăng nhập không chính xác".',
    '5. bcrypt.compare(password, user.passwordHash) so sánh.',
    '6. Sai mật khẩu → 401.',
    '7. Đúng → jwt.sign({id,role,username}, JWT_SECRET, {expiresIn:"7d"}).',
    '8. Trả 200 {token, user: sanitizeUser(user)} (đã loại bỏ passwordHash).',
]: N(doc, s)

H(doc, '4.3 Sơ đồ Trình tự – Thêm cư dân & Tạo tài khoản', 2)
for s in [
    '1. Client → POST /api/residents {name, cccd, dob, gender, room, regdate, ...}',
    '2. requireAuth: jwt.verify() + User.findById() → lỗi trả 401.',
    '3. requireRole("admin","staff"): role không khớp → trả 403.',
    '4. Kiểm tra trường bắt buộc (cccd, name, dob, gender, room, regdate) → thiếu trả 400.',
    '5. Resident.findOne({cccd}) → CCCD đã tồn tại → trả 409 Conflict.',
    '6. Resident.create({...data, status:"Thường trú", history:[{action:"Thêm mới cư dân"}]}).',
    '7. defaultPassword = cccd.slice(-8) (8 số cuối CCCD).',
    '8. bcrypt.hash(defaultPassword, 10) mã hóa.',
    '9. User.create({username:cccd, email:cccd+"@resident.local", role:"resident", mustChangePassword:true, residentId}).',
    '10. Trả 201 Created {message, resident}.',
]: N(doc, s)

H(doc, '4.4 Sơ đồ Trình tự – Đăng ký Tạm trú', 2)
for s in [
    '1. Client → POST /api/residents/:id/tam-tru {address, start, end, reason, phone}.',
    '2. requireAuth + requireRole("admin","staff").',
    '3. Resident.findById(id) → không tìm thấy → 404.',
    '4. Thiếu address/start/end/phone → 400.',
    '5. end < start → 400 "Ngày kết thúc không được nhỏ hơn ngày bắt đầu".',
    '6. status == "Tạm vắng" → 409 "Đang tạm vắng, không thể đăng ký tạm trú".',
    '7. resident.status = "Tạm trú"; resident.tamTru = {address, start, end, reason, phone}.',
    '8. addHistory(resident, "Đăng ký tạm trú", req.user.name) → push vào history[].',
    '9. resident.save() → lưu MongoDB Atlas.',
    '10. Trả 200 {message, resident}.',
]: N(doc, s)

H(doc, '4.5 Frontend – Cấu trúc màn hình', 2)
T(doc, ['Route', 'Component', 'Quyền', 'Chức năng'],
[['/dashboard', 'Dashboard.jsx', 'admin, staff', 'Thống kê: cards số lượng, Line Chart 6 tháng (Chart.js), log 10 hoạt động gần nhất'],
 ['/residents', 'ResidentList.jsx', 'admin, staff', 'Danh sách cư dân, tìm kiếm realtime, lọc theo trạng thái'],
 ['/residents/add', 'ResidentAdd.jsx', 'admin, staff', 'Form thêm cư dân mới'],
 ['/residents/:id', 'ResidentDetail.jsx', 'admin, staff', 'Xem/sửa hồ sơ, đăng ký tạm trú/vắng, lịch sử biến động'],
 ['/tamtru', 'TamTru.jsx', 'admin, staff', 'Form đăng ký Tạm trú'],
 ['/tamvang', 'TamVang.jsx', 'admin, staff', 'Form đăng ký Tạm vắng'],
 ['/reports', 'Reports.jsx', 'admin, staff', 'Báo cáo thống kê'],
 ['/my-profile', 'MyProfile.jsx', 'resident', 'Cư dân xem hồ sơ cá nhân (GET /api/residents/me)']])

# ========== CHƯƠNG 5 ==========
H(doc, 'CHƯƠNG 5. TRIỂN KHAI VÀ KIỂM THỬ (Tuần 7)', 1)

H(doc, '5.1 Thuật toán Business Rules', 2)
P(doc, 'BR-006 – Không xóa cư dân đang lưu trú:', bold=True)
for s in [
    'Tìm cư dân theo id → không thấy: 404.',
    'status == "Tạm trú" → 409 "Không thể xóa cư dân đang tạm trú".',
    'status == "Tạm vắng" → 409 "Không thể xóa cư dân đang tạm vắng".',
    'User.deleteOne({residentId}) xóa tài khoản liên kết.',
    'resident.deleteOne() → 200 "Xóa cư dân thành công".',
]: B(doc, s)

P(doc, 'Thuật toán getResidentStats (Dashboard):', bold=True)
for s in [
    'Resident.find() lấy toàn bộ (chỉ field: status, createdAt, history, name).',
    'Đếm total, thuongtru, tamtru, tamvang bằng Array.filter() theo status.',
    'Vòng lặp 6 tháng gần nhất: mỗi tháng đếm cư dân có createdAt <= endOfMonth.',
    'Gộp tất cả history[], sort giảm dần, lấy 10 mục → recentActivities.',
    'Trả JSON: {total, thuongtru, tamtru, tamvang, monthlyStats[6], recentActivities[10]}.',
]: B(doc, s)

H(doc, '5.2 Xử lý lỗi', 2)
T(doc, ['HTTP Code', 'Trường hợp trong source code'],
[['400', 'Thiếu trường bắt buộc. end < start (tạm trú/vắng). Mật khẩu < 8 ký tự.'],
 ['401', 'Sai username/password. Không có Bearer token. Token hết hạn (7 ngày).'],
 ['403', 'Resident gọi API admin/staff. Staff gọi API chỉ dành cho admin (DELETE).'],
 ['404', 'id cư dân không tồn tại. residentId null khi gọi /residents/me.'],
 ['409', 'CCCD trùng. Đăng ký tạm trú khi đang tạm vắng (và ngược lại). Xóa cư dân đang lưu trú.'],
 ['500', 'Mất kết nối MongoDB Atlas. Lỗi ngoại lệ không bắt được.']])

H(doc, '5.3 Kiểm thử', 2)
T(doc, ['Loại', 'Công cụ', 'Kịch bản'],
[['Integration (API)', 'Postman', 'Test 14 endpoints. Negative: CCCD trùng → 409, tạm trú khi đang tạm vắng → 409, xóa khi đang lưu trú → 409, không token → 401, sai role → 403.'],
 ['Manual (UI)', 'Chrome Browser', 'Đăng nhập 3 role, kiểm tra phân quyền menu. Thêm cư dân → đăng nhập bằng CCCD (mật khẩu 8 số cuối) → bắt buộc đổi mật khẩu. Dashboard cập nhật số liệu sau thao tác.']])

H(doc, '5.4 Tiêu chuẩn Code & Kế hoạch', 2)
T(doc, ['Quy tắc', 'Áp dụng'],
[['Đặt tên', 'camelCase cho hàm/biến (getResidentStats, addHistory). PascalCase cho Model (Resident, User).'],
 ['File', '<module>.<type>.js: auth.controller.js, resident.routes.js, auth.middleware.js'],
 ['Module System', 'ES Modules (import/export) cho cả Backend và Frontend (package.json "type":"module")'],
 ['Error format', 'Mọi lỗi trả {message: "..."}. sanitizeUser() loại bỏ passwordHash khỏi response.']])

H(doc, 'KẾT LUẬN', 1)
P(doc, 'Dự án ResidentIQ đã được thiết kế và triển khai hoàn chỉnh bằng MERN Stack (MongoDB Atlas + Express.js + React.js 19 + Node.js). Hệ thống đáp ứng đầy đủ 21 yêu cầu chức năng và 15 yêu cầu phi chức năng. Điểm nổi bật: kiến trúc Modular Monolith phù hợp team 3 người; phân quyền 3 cấp bảo vệ ở cả Backend (middleware) lẫn Frontend (route guard); tự động tạo tài khoản User khi thêm cư dân; lịch sử biến động ghi tự động; Business Rules BR-001 (CCCD unique), BR-005 (không đồng thời tạm trú+tạm vắng), BR-006 (không xóa khi đang lưu trú) được cài đặt chặt chẽ trong ResidentController.')

doc.save(r'd:\TTCS-main\BaoCaoCuoiKy_Fixed.docx')
print('DONE')
